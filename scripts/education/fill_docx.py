#!/usr/bin/env python3
"""
fill_docx.py — 依據 LLM 產出的 JSON 指令，定點填入 .docx 模板。
保留原始格式，只修改指定位置的文字。

Usage:
    python3 fill_docx.py <template.docx> <fill_instructions.json> <output.docx>

fill_instructions.json 格式：
{
  "fills": [
    {
      "type": "paragraph",
      "index": 5,
      "action": "replace",
      "text": "新內容"
    },
    {
      "type": "table",
      "table_index": 0,
      "row": 2,
      "col": 1,
      "action": "replace",
      "text": "填入內容"
    },
    {
      "type": "paragraph",
      "index": 10,
      "action": "append",
      "text": "追加的內容"
    },
    {
      "type": "checkbox",
      "index": 15,
      "checks": ["□民主"],
      "description": "勾選教養態度"
    }
  ]
}
"""

import json
import sys
import os
import copy

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("請安裝 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def preserve_format_replace(paragraph, new_text: str):
    """替換段落文字，保留第一個 run 的格式。"""
    if not paragraph.runs:
        run = paragraph.add_run(new_text)
        return

    # 保留第一個 run 的格式
    first_run = paragraph.runs[0]
    fmt = {
        "bold": first_run.font.bold,
        "italic": first_run.font.italic,
        "size": first_run.font.size,
        "name": first_run.font.name,
    }

    # 取得東亞字型
    r_elem = first_run._element
    rPr = r_elem.find(qn("w:rPr"))
    east_asia_font = None
    if rPr is not None:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            east_asia_font = rFonts.get(qn("w:eastAsia"))

    # 清除所有 run
    for run in paragraph.runs:
        run.clear()

    # 只保留第一個 run，寫入新文字
    paragraph.runs[0].text = new_text
    paragraph.runs[0].font.bold = fmt["bold"]
    paragraph.runs[0].font.italic = fmt["italic"]
    if fmt["size"]:
        paragraph.runs[0].font.size = fmt["size"]
    if fmt["name"]:
        paragraph.runs[0].font.name = fmt["name"]

    # 還原東亞字型
    if east_asia_font:
        r = paragraph.runs[0]._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), east_asia_font)

    # 刪除多餘的 run
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)


def handle_checkbox(paragraph, checks: list):
    """將 □ 替換為 ☑（勾選）。"""
    full_text = paragraph.text
    for check_item in checks:
        # 把 □ 開頭的選項改成 ☑
        unchecked = check_item
        checked = check_item.replace("□", "☑")
        full_text = full_text.replace(unchecked, checked)

    # 重寫整段
    preserve_format_replace(paragraph, full_text)


def fill_docx(template_path: str, instructions: dict, output_path: str):
    """依據指令填入 .docx 模板。"""
    doc = Document(template_path)

    for fill in instructions.get("fills", []):
        fill_type = fill.get("type")
        action = fill.get("action", "replace")

        try:
            if fill_type == "paragraph":
                idx = fill["index"]
                if idx < len(doc.paragraphs):
                    para = doc.paragraphs[idx]
                    if action == "replace":
                        preserve_format_replace(para, fill["text"])
                    elif action == "append":
                        para.add_run("\n" + fill["text"])

            elif fill_type == "table":
                ti = fill["table_index"]
                row = fill["row"]
                col = fill["col"]
                if ti < len(doc.tables):
                    table = doc.tables[ti]
                    if row < len(table.rows) and col < len(table.rows[row].cells):
                        cell = table.rows[row].cells[col]
                        if action == "replace":
                            # 保留 cell 格式，只改文字
                            if cell.paragraphs:
                                preserve_format_replace(cell.paragraphs[0], fill["text"])
                            else:
                                cell.text = fill["text"]
                        elif action == "append":
                            if cell.paragraphs:
                                cell.paragraphs[0].add_run("\n" + fill["text"])

            elif fill_type == "checkbox":
                idx = fill["index"]
                if idx < len(doc.paragraphs):
                    handle_checkbox(doc.paragraphs[idx], fill.get("checks", []))

        except Exception as e:
            print(f"Warning: Failed to fill {fill_type} at index {fill.get('index', fill.get('table_index', '?'))}: {e}", file=sys.stderr)
            continue

    doc.save(output_path)
    print(f"Document saved to: {output_path}")


def main():
    if len(sys.argv) < 4:
        print(f"Usage: {sys.argv[0]} <template.docx> <fill_instructions.json> <output.docx>", file=sys.stderr)
        sys.exit(1)

    template_path = sys.argv[1]
    instructions_path = sys.argv[2]
    output_path = sys.argv[3]

    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}", file=sys.stderr)
        sys.exit(1)

    with open(instructions_path, "r", encoding="utf-8") as f:
        instructions = json.load(f)

    fill_docx(template_path, instructions, output_path)


if __name__ == "__main__":
    main()
