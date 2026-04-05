"""
docx_utils.py — 共用 python-docx 工具函數與格式常數

供 IEP 會議記錄、特推會記錄等文件生成腳本使用。
"""

from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 格式常數 ──

CHINESE_FONT = "標楷體"
ENGLISH_FONT = "Times New Roman"
TITLE_SIZE = Pt(14)
BODY_SIZE = Pt(10)
SMALL_SIZE = Pt(10)

PAGE_WIDTH = Emu(7772400)   # A4
PAGE_HEIGHT = Emu(10058400)


# ── 工具函數 ──

def setup_page(section):
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_run_font(run, size=None, bold=None):
    if size is None:
        size = BODY_SIZE
    run.font.size = size
    run.font.bold = bold
    run.font.name = ENGLISH_FONT
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), CHINESE_FONT)


def add_paragraph(doc, text, size=None, bold=None, alignment=None,
                  space_before=None, space_after=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    fmt = p.paragraph_format
    fmt.space_before = space_before or Pt(0)
    fmt.space_after = space_after or Pt(0)
    # 行距設為單行
    fmt.line_spacing = Pt(18) if size and size <= Pt(12) else Pt(22)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def set_cell_text(cell, text, size=None, bold=None,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, vertical_center=True):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size or BODY_SIZE, bold=bold)
    if vertical_center:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        tcPr.append(vAlign)


def set_cell_multiline(cell, text, size=None):
    """多行文字寫入 cell，保留換行。"""
    cell.text = ""
    for i, line in enumerate(text.split("\n")):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = Pt(16)
        run = p.add_run(line)
        set_run_font(run, size=size or BODY_SIZE)


def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    tblPr.append(borders)


def merge_cells(table, row, col_start, col_end):
    """合併同一行的多個欄位。"""
    cell_start = table.cell(row, col_start)
    cell_end = table.cell(row, col_end)
    cell_start.merge(cell_end)
