#!/usr/bin/env python3
"""
iep_meeting_generator.py — 從零生成 IEP 會議記錄 .docx

根據 LLM 分析結果（JSON），生成一份乾淨的 IEP 會議記錄文件。
格式參考南投縣 IEP 會議記錄模板（113/114 學年度格式）。

Usage:
    python3 iep_meeting_generator.py <content.json> --output output.docx
"""

import json
import sys
import os
import argparse

from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

from docx_utils import (
    CHINESE_FONT, ENGLISH_FONT, TITLE_SIZE, BODY_SIZE, SMALL_SIZE,
    PAGE_WIDTH, PAGE_HEIGHT,
    setup_page, set_run_font, add_paragraph,
    set_cell_text, set_cell_multiline, add_table_borders, merge_cells,
)


# ── 格式常數（本檔專用）──

SIGN_SIZE = Pt(12)


# ── 輔助函數 ──

def set_col_width(table, col_idx, width):
    for row in table.rows:
        row.cells[col_idx].width = width


# ── 會議記錄生成 ──

def generate_meeting_record(data: dict, output_path: str):
    """
    data 格式：
    {
        "school_name": "○○",
        "academic_year": "114",
        "semester": 1,
        "meeting_type": "擬訂" | "檢討",
        "student_name": "王小明",
        "meeting_date": "114年6月26日",
        "location": "2F會議室",
        "chair": "○○○",
        "recorder": "○○○",
        "discussion": "討論內容...",
        "resolution": "決議內容...",
        "attendees": {
            "admin": ["○○○"],
            "parents": ["○○○"],
            "regular_teachers": ["○○○"],
            "special_ed_teachers": ["○○○", "○○○"],
            "professionals": [],
            "assistants": ["○○○"]
        }
    }
    """
    doc = Document()
    setup_page(doc.sections[0])

    school = data.get("school_name", "____")
    year = data.get("academic_year", "___")
    semester = data.get("semester", 1)
    mtype = data.get("meeting_type", "擬訂")
    student = data.get("student_name", "________")
    mdate = data.get("meeting_date", "   年   月   日")
    location = data.get("location", "________")
    chair = data.get("chair", "________")
    recorder = data.get("recorder", "________")

    # ── 標題 ──
    add_paragraph(
        doc,
        f"南投縣 {school} 國小 {year} 學年度第 {semester} 學期學生個別化教育計畫{mtype}會議",
        size=TITLE_SIZE, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(6),
    )

    # ── 基本資訊 ──
    add_paragraph(doc, f"學生姓名：{student}　　　　　　　　　　　　　　　會議日期：{mdate}", size=BODY_SIZE)
    add_paragraph(doc, f"地　　點：{location}", size=BODY_SIZE)
    add_paragraph(doc, f"主　　席：{chair}　　　　　　　　　　　　　　　　記錄者：{recorder}", size=BODY_SIZE,
                  space_after=Pt(6))

    # ── 簽到表 ──
    attendees = data.get("attendees", {})
    cols = ["行政人員", "學生家長\n及學生", "普通班教師", "特教教師", "相關專業人員", "相關助理人員"]
    att_data = [
        attendees.get("admin", []),
        attendees.get("parents", []),
        attendees.get("regular_teachers", []),
        attendees.get("special_ed_teachers", []),
        attendees.get("professionals", []),
        attendees.get("assistants", []),
    ]

    max_rows = 3  # 固定 3 行簽名欄

    sign_table = doc.add_table(rows=2 + max_rows, cols=6)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 第一行：合併為「與會者簽名」
    merge_cells(sign_table, 0, 0, 5)
    set_cell_text(sign_table.cell(0, 0), "與會者簽名", size=BODY_SIZE, bold=True)

    # 第二行：欄位標題
    for ci, header in enumerate(cols):
        set_cell_text(sign_table.cell(1, ci), header, size=SMALL_SIZE, bold=True)

    # 簽名行：留空（由與會人員現場簽名）
    for ri in range(max_rows):
        for ci in range(6):
            set_cell_text(sign_table.cell(2 + ri, ci), "", size=BODY_SIZE)

    add_table_borders(sign_table)

    # 間距
    add_paragraph(doc, "", space_after=Pt(4))

    # ── 會議紀錄表 ──
    record_table = doc.add_table(rows=5, cols=1)
    record_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: 標題
    set_cell_text(record_table.cell(0, 0), "會議紀錄", size=BODY_SIZE, bold=True)

    # Row 1: 子標題
    set_cell_text(record_table.cell(1, 0), "個案報告及討論事項", size=BODY_SIZE, bold=True)

    # Row 2: 討論內容（主要內容）
    discussion = data.get("discussion", "")
    set_cell_multiline(record_table.cell(2, 0), discussion, size=BODY_SIZE)

    # Row 3: 決議標題
    set_cell_text(record_table.cell(3, 0), "會議決議", size=BODY_SIZE, bold=True)

    # Row 4: 決議內容
    resolution = data.get("resolution", "")
    set_cell_multiline(record_table.cell(4, 0), resolution, size=BODY_SIZE)

    add_table_borders(record_table)

    # ── 簽名欄 ──
    add_paragraph(doc, "", space_after=Pt(8))

    if mtype == "擬訂":
        add_paragraph(doc, "本人已詳細閱讀，並同意這份計畫的執行！", size=BODY_SIZE)
        add_paragraph(doc, f"日期：　　　　　　　　　　　　　　家長簽名：____________________", size=BODY_SIZE,
                      space_after=Pt(12))
    else:
        add_paragraph(doc, f"家長確認簽名：____________________　　　　　日期：", size=BODY_SIZE,
                      space_after=Pt(12))

    add_paragraph(doc, "承辦人：　　　　　　　　主管：　　　　　　　　校長：", size=BODY_SIZE)

    doc.save(output_path)
    return output_path


# ── CLI ──

def main():
    parser = argparse.ArgumentParser(description="從零生成 IEP 會議記錄")
    parser.add_argument("content_json", help="LLM 產出的內容 JSON")
    parser.add_argument("--output", required=True, help="輸出 .docx 路徑")

    args = parser.parse_args()

    with open(args.content_json, "r", encoding="utf-8") as f:
        data = json.load(f)

    output = generate_meeting_record(data, args.output)
    print(f"✅ 已生成：{output}")


if __name__ == "__main__":
    main()
