#!/usr/bin/env python3
"""
spc_meeting_core.py — 特推會會議記錄產生器核心邏輯

供 CLI 和未來 Dashboard API 共用。
"""

import glob
import json
import os
import re
import subprocess
import sys
from dataclasses import dataclass, field

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from docx_utils import (
    setup_page, set_run_font, add_paragraph,
    set_cell_text, set_cell_multiline, add_table_borders, merge_cells,
    TITLE_SIZE, BODY_SIZE,
)


# ── 路徑常數 ──

OBSIDIAN_SPC_DIR = os.path.expanduser(
    "~/Library/CloudStorage/GoogleDrive-user@gmail.com"
    "/我的雲端硬碟/Obsidian-Cyclone/02-特教業務/特推會"
)


# ── 學校預設值與委員名單 ──

SCHOOL_DEFAULTS = {
    "school_name": "○○",
    "chair": "○○○",
    "recorder": "○○○",
    "location": "本校三樓共讀站",
    "academic_year": 114,
}

DEFAULT_COMMITTEE = [
    {"title": "主任委員", "role": "校長", "name": "○○○"},
    {"title": "副主任委員\n兼主任秘書", "role": "教導主任", "name": "○○○"},
    {"title": "副主任委員", "role": "總務主任", "name": "○○○"},
    {"title": "委員 兼\n執行秘書", "role": "特教教師", "name": "○○○"},
    {"title": "委員", "role": "教務組長", "name": "○○○"},
    {"title": "委員", "role": "特教教師", "name": "○○○"},
    {"title": "委員", "role": "普通班教師", "name": "○○○"},
    {"title": "委員", "role": "普通班教師", "name": "○○○"},
    {"title": "委員", "role": "家長代表", "name": "○○○"},
]

PROPOSAL_TYPES = [
    "交通補助",
    "專團申請",
    "助理員申請",
    "酌減學生數",
    "轉安置",
    "課程計畫審議",
    "鑑定安置提報",
    "其他",
]


# ── 資料類別 ──

@dataclass
class Proposal:
    type: str
    title: str
    description: str = ""
    decision: str = "（會後填入）"
    students: list = field(default_factory=list)
    ref_doc: str = ""


@dataclass
class MeetingRecord:
    academic_year: int = 114
    meeting_number: int = 1
    date: str = ""
    weekday: str = ""
    time_start: str = "上午08:10"
    time_end: str = ""
    location: str = "本校三樓共讀站"
    chair: str = "○○○"
    recorder: str = "○○○"
    business_report: str = ""
    previous_tracking: str = ""
    proposals: list = field(default_factory=list)
    motions: str = "無"
    committee: list = field(default_factory=list)


# ── 關鍵字對應 ──

PROPOSAL_KEYWORDS = {
    "交通補助": ["交通補助"],
    "專團申請": ["專團", "專業人員"],
    "助理員申請": ["助理員", "教助"],
    "酌減學生數": ["酌減"],
    "轉安置": ["轉安置"],
    "課程計畫審議": ["課程計畫", "審議"],
    "鑑定安置提報": ["鑑定安置", "提報"],
}


# ── 歷史資料讀取 ──

def fetch_similar_meetings(proposal_type, n=3):
    """
    從 Obsidian 資料夾讀取歷次特推會記錄，找出與指定提案類型相近的案例。

    Returns:
        list of dicts: [{filename, academic_year, meeting_number, descriptions}]
    """
    keywords = PROPOSAL_KEYWORDS.get(proposal_type, [proposal_type])
    pattern = os.path.join(OBSIDIAN_SPC_DIR, "*.md")
    files = sorted(glob.glob(pattern), reverse=True)

    results = []
    for filepath in files:
        filename = os.path.basename(filepath)
        if filename.startswith("MOC-"):
            continue

        try:
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
        except OSError:
            continue

        # 檢查 frontmatter 的 topics 或檔名是否符合關鍵字
        matched = False
        for kw in keywords:
            if kw in filename or kw in content[:500]:
                matched = True
                break

        if not matched:
            continue

        # 提取說明段落
        desc_blocks = re.findall(
            r"\*\*說明：\*\*\s*\n\n(.+?)(?=\n\*\*決議|\n###|\Z)",
            content,
            re.DOTALL,
        )
        descriptions = [d.strip() for d in desc_blocks if d.strip()]

        # 從 frontmatter 提取 academic_year 和 meeting_number
        academic_year = None
        meeting_number = None
        fm_match = re.search(r"^---\n(.+?)\n---", content, re.DOTALL)
        if fm_match:
            fm = fm_match.group(1)
            yr_m = re.search(r"academic_year:\s*(\d+)", fm)
            num_m = re.search(r"meeting_number:\s*(\d+)", fm)
            if yr_m:
                academic_year = int(yr_m.group(1))
            if num_m:
                meeting_number = int(num_m.group(1))

        results.append({
            "filename": filename,
            "academic_year": academic_year,
            "meeting_number": meeting_number,
            "descriptions": descriptions,
        })

        if len(results) >= n:
            break

    return results


def fetch_previous_decisions(academic_year, meeting_number):
    """
    讀取上一次特推會的決議事項，供前次決議追蹤使用。

    Returns:
        str: 格式化的前次決議文字，若無則回傳空字串。
    """
    if meeting_number <= 1:
        return ""

    prev_num = meeting_number - 1
    pattern = os.path.join(
        OBSIDIAN_SPC_DIR,
        f"{academic_year}-特推會-{prev_num:02d}-*.md",
    )
    files = glob.glob(pattern)
    if not files:
        return ""

    filepath = files[0]
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
    except OSError:
        return ""

    # 提取決議段落
    decision_blocks = re.findall(
        r"\*\*決議：\*\*\s*\n\n(.+?)(?=\n###|\n##|\Z)",
        content,
        re.DOTALL,
    )

    if not decision_blocks:
        return ""

    # 取得主題
    topic = ""
    fm_match = re.search(r"^---\n(.+?)\n---", content, re.DOTALL)
    if fm_match:
        fm = fm_match.group(1)
        topics_m = re.search(r"topics:\s*\n((?:\s+-.+\n?)+)", fm)
        if topics_m:
            topic_lines = re.findall(r"-\s*(.+)", topics_m.group(1))
            topic = "、".join(t.strip() for t in topic_lines)

    lines = [f"前次會議（第{prev_num}次，{topic}）決議事項："]
    for i, block in enumerate(decision_blocks, 1):
        stripped = block.strip()
        lines.append(f"{i}. {stripped}")

    return "\n".join(lines)


# ── Markdown 輸出 ──

def build_markdown(record: MeetingRecord, status: str = "record-generated", mode: str = "record") -> str:
    """生成符合 Obsidian 特推會格式的 Markdown 文件。"""

    # 蒐集 topics 和 decisions
    topics = [p.type for p in record.proposals]
    decisions = [p.decision for p in record.proposals]

    # YAML frontmatter
    topics_yaml = "\n".join(f"  - {t}" for t in topics) if topics else "  - 未分類"
    decisions_yaml = "\n".join(f'  - "{d}"' for d in decisions) if decisions else ""

    frontmatter_lines = [
        "---",
        "type: 特推會會議記錄",
        f"academic_year: {record.academic_year}",
        f"meeting_number: {record.meeting_number}",
        f'date: "{record.date}"',
        f'chair: "{record.chair}"',
        f'recorder: "{record.recorder}"',
        f'location: "{record.location}"',
        "topics:",
        topics_yaml,
    ]
    if decisions_yaml:
        frontmatter_lines.append("decisions:")
        frontmatter_lines.append(decisions_yaml)
    else:
        frontmatter_lines.append("decisions:")
    frontmatter_lines.append("tags: [特推會, 會議記錄]")
    frontmatter_lines.append(f'status: "{status}"')
    frontmatter_lines.append(f'mode: "{mode}"')
    frontmatter_lines.append("---")

    frontmatter = "\n".join(frontmatter_lines)

    # 文件主題（多個用 & 連接）
    topics_str = "&".join(topics) if topics else "未分類"

    # 標題與基本資訊
    body_lines = [
        f"# {record.academic_year}學年度 第{record.meeting_number}次特推會會議記錄",
        "",
        f"- **日期**：{record.date}",
        f"- **地點**：{record.location}",
        f"- **主席**：{record.chair}",
        f"- **記錄**：{record.recorder}",
        f"- **主題**：{topics_str}",
        "",
        "---",
        "",
    ]

    # 業務報告
    body_lines.append("## 業務報告")
    body_lines.append("")
    if record.business_report:
        body_lines.append(record.business_report)
    body_lines.append("")

    # 前次會議決議追蹤
    if record.previous_tracking:
        body_lines.append("## 前次會議決議追蹤")
        body_lines.append("")
        body_lines.append(record.previous_tracking)
        body_lines.append("")

    # 提案討論
    if record.proposals:
        body_lines.append("## 提案討論")
        body_lines.append("")
        for i, p in enumerate(record.proposals, 1):
            body_lines.append(f"### 案由{i}：{p.title}")
            body_lines.append("")
            body_lines.append("**說明：**")
            body_lines.append("")
            body_lines.append(p.description if p.description else "（待填）")
            body_lines.append("")
            body_lines.append("**決議：**")
            body_lines.append("")
            body_lines.append(p.decision)
            body_lines.append("")

    # 臨時動議
    body_lines.append("## 臨時動議")
    body_lines.append("")
    body_lines.append(record.motions if record.motions else "無")
    body_lines.append("")

    # 散會
    body_lines.append("## 散會")
    body_lines.append("")
    if record.time_end:
        body_lines.append(record.time_end)
    else:
        body_lines.append("（散會時間）")
    body_lines.append("")

    body = "\n".join(body_lines)
    return f"{frontmatter}\n\n{body}"


# ── 儲存 Markdown ──

def save_markdown(record: MeetingRecord, status: str = "record-generated", mode: str = "record") -> str:
    """儲存 Markdown 至 Obsidian 特推會資料夾，並更新 MOC。"""
    os.makedirs(OBSIDIAN_SPC_DIR, exist_ok=True)

    topics = [p.type for p in record.proposals]
    topics_str = "+".join(topics) if topics else "未分類"
    filename = f"{record.academic_year}-特推會-{record.meeting_number:02d}-{topics_str}.md"
    filepath = os.path.join(OBSIDIAN_SPC_DIR, filename)

    content = build_markdown(record, status=status, mode=mode)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

    update_moc()
    return filepath


# ── 更新 MOC ──

def update_moc():
    """讀取所有特推會 .md 檔案，重新生成 MOC-特推會.md。"""
    pattern = os.path.join(OBSIDIAN_SPC_DIR, "*.md")
    files = sorted(glob.glob(pattern))

    # 依學年度分組
    by_year = {}
    for filepath in files:
        filename = os.path.basename(filepath)
        if filename.startswith("MOC-"):
            continue

        try:
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
        except OSError:
            continue

        fm_match = re.search(r"^---\n(.+?)\n---", content, re.DOTALL)
        if not fm_match:
            continue

        fm = fm_match.group(1)
        yr_m = re.search(r"academic_year:\s*(\d+)", fm)
        num_m = re.search(r"meeting_number:\s*(\d+)", fm)
        date_m = re.search(r'date:\s*"?([^"\n]+)"?', fm)
        topics_m = re.search(r"topics:\s*\n((?:\s+-.+\n?)+)", fm)
        decisions_m = re.search(r"decisions:\s*\n((?:\s+-.+\n?)+)", fm)

        if not yr_m or not num_m:
            continue

        year = int(yr_m.group(1))
        num = int(num_m.group(1))
        date = date_m.group(1).strip() if date_m else ""
        topics_list = re.findall(r"-\s*(.+)", topics_m.group(1)) if topics_m else []
        topics_str = "&".join(t.strip() for t in topics_list)
        decisions_list = re.findall(r'-\s*"?([^"\n]+)"?', decisions_m.group(1)) if decisions_m else []
        # 決議摘要截斷至 30 字
        decisions_summary = "; ".join(
            d.strip()[:30] for d in decisions_list if d.strip()
        )

        stem = filename.replace(".md", "")

        if year not in by_year:
            by_year[year] = []
        by_year[year].append({
            "num": num,
            "date": date,
            "topics": topics_str,
            "decisions": decisions_summary,
            "stem": stem,
        })

    # 生成 MOC 內容
    lines = [
        "---",
        "type: MOC",
        "tags: [特推會, MOC]",
        "---",
        "",
        "# 特推會會議記錄索引",
        "",
        "> ○○國小特殊教育推行委員會歷次會議記錄",
        "",
    ]

    for year in sorted(by_year.keys()):
        lines.append(f"## {year} 學年度")
        lines.append("")
        lines.append("| 次 | 日期 | 主題 | 決議摘要 |")
        lines.append("|---|---|---|---|")
        meetings = sorted(by_year[year], key=lambda x: x["num"])
        for m in meetings:
            link = f"[[{m['stem']}|第{m['num']}次]]"
            lines.append(f"| {link} | {m['date']} | {m['topics']} | {m['decisions']} |")
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("*由 CycloneOS 自動生成*")

    moc_path = os.path.join(OBSIDIAN_SPC_DIR, "MOC-特推會.md")
    with open(moc_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


# ── LLM 呼叫 ──

def call_llm(prompt: str) -> str:
    """呼叫本機 LLM CLI（優先 claude，備選 codex）。"""
    env = {k: v for k, v in os.environ.items() if k != "ANTHROPIC_API_KEY"}

    candidates = [
        ["claude", "-p", prompt],
        ["codex", "--quiet", "--full-auto", prompt],
    ]

    for cmd in candidates:
        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=300,
                env=env,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    raise RuntimeError("所有 LLM 工具皆無法使用（claude / codex）")


# ── LLM 起草提案說明 ──

def draft_proposal(proposal: Proposal, similar: list) -> str:
    """使用 LLM 起草提案的說明段落。"""

    # 整理學生資訊
    student_info = ""
    if proposal.students:
        parts = []
        for s in proposal.students:
            name = mask_name(s.get("name", ""))
            grade = s.get("grade", "")
            disability = s.get("disability", "")
            detail = s.get("detail", "")
            parts.append(f"  - {grade} {name}：{disability}。{detail}")
        student_info = "\n".join(parts)
    else:
        student_info = "  （無學生資訊）"

    # 整理歷史參考
    history_parts = []
    for item in similar:
        fname = item.get("filename", "")
        yr = item.get("academic_year", "")
        num = item.get("meeting_number", "")
        descs = item.get("descriptions", [])
        if descs:
            history_parts.append(f"【{yr}學年度第{num}次，{fname}】")
            for d in descs[:2]:
                history_parts.append(d[:300])

    history_text = "\n\n".join(history_parts) if history_parts else "（無歷史參考）"

    prompt = f"""你是台灣特殊教育行政人員，正在撰寫特推會（特殊教育推行委員會）會議記錄中的提案說明段落。

## 案由資訊
- 案由類型：{proposal.type}
- 案由標題：{proposal.title}
- 公文字號：{proposal.ref_doc if proposal.ref_doc else "（無）"}

## 學生資訊
{student_info}

## 歷史參考說明段落
{history_text}

## 撰寫原則
1. 使用正式公文用語，語氣嚴謹
2. 參考歷次格式，保持一致性
3. 學生姓名中間字用○代替（例：王小明→王○明）
4. 如有公文字號，請在說明中引用
5. 使用繁體中文
6. 段落用（一）（二）編號
7. 直接輸出說明段落本文，不要加標題或前言

請撰寫此案由的說明段落："""

    return call_llm(prompt)


# ── 姓名遮蔽 ──

def mask_name(name: str) -> str:
    """遮蔽姓名中間字。"""
    if len(name) <= 1:
        return name
    if len(name) == 2:
        return name[0] + "○"
    if len(name) == 3:
        return name[0] + "○" + name[2]
    # 4+ chars: keep first and last
    return name[0] + "○" * (len(name) - 2) + name[-1]


_ADDRESS_RE = re.compile(
    r"[\u4e00-\u9fff]{1,3}[縣市][\u4e00-\u9fff]{1,4}[鄉鎮市區][\u4e00-\u9fff\d\-號巷弄樓之]+"
)
_PHONE_RE = re.compile(
    r"(?:0\d{1,2}[-\s]?\d{3,4}[-\s]?\d{3,4}|09\d{2}[-\s]?\d{3}[-\s]?\d{3})"
)


def mask_pii(text: str, names: list[str]) -> str:
    """Mask all PII in text: known names, addresses, phone numbers."""
    result = text
    for name in sorted(names, key=len, reverse=True):
        if name and len(name) >= 2:
            result = result.replace(name, mask_name(name))
    result = _ADDRESS_RE.sub("（地址已隱藏）", result)
    result = _PHONE_RE.sub("（電話已隱藏）", result)
    return result


def collect_names(record) -> list[str]:
    """Collect all names from a MeetingRecord for PII masking."""
    names = set()
    names.add(record.chair)
    names.add(record.recorder)
    for m in record.committee:
        if isinstance(m, dict):
            names.add(m.get("name", ""))
        elif hasattr(m, "name"):
            names.add(m.name)
    for p in record.proposals:
        for s in (p.students if p.students else []):
            if isinstance(s, dict):
                names.add(s.get("name", ""))
            elif isinstance(s, str):
                names.add(s)
    names.discard("")
    return list(names)


# ── 生成 .docx ──

CHINESE_NUMS = "一二三四五六七八九十"


def _chinese_num(n: int) -> str:
    """將數字轉換為中文數字（1-10）。"""
    if 1 <= n <= len(CHINESE_NUMS):
        return CHINESE_NUMS[n - 1]
    return str(n)


def generate_docx(record: MeetingRecord, output_path: str, mask_names: list[str] | None = None) -> str:
    """從零生成特推會會議記錄 .docx。"""

    def m(text):
        """Apply PII masking if names provided."""
        if mask_names and text:
            return mask_pii(text, mask_names)
        return text if text else ""

    doc = Document()
    setup_page(doc.sections[0])

    school = SCHOOL_DEFAULTS["school_name"]
    year = record.academic_year
    num = record.meeting_number
    committee = record.committee if record.committee else DEFAULT_COMMITTEE

    # ════════════════════════════════════════
    # 第一頁：會議記錄本文
    # ════════════════════════════════════════

    # ── 標題 ──
    add_paragraph(
        doc,
        f"南投縣{school}國民小學{year}學年度特殊教育推行委員會",
        size=TITLE_SIZE, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(2),
    )
    add_paragraph(
        doc,
        f"第{num}次會議紀錄",
        size=TITLE_SIZE, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(6),
    )

    # ── 壹、會議時間 ──
    time_str = record.time_start
    if record.date and record.weekday:
        time_str = f"{record.date}（{record.weekday}）{record.time_start}"
    elif record.date:
        time_str = f"{record.date} {record.time_start}"
    add_paragraph(doc, f"壹、會議時間：{time_str}", size=BODY_SIZE)

    # ── 貳、會議地點 ──
    add_paragraph(doc, f"貳、會議地點：{record.location}", size=BODY_SIZE)

    # ── 參、主席 + 記錄 ──
    add_paragraph(
        doc,
        f"參、主　　席：{m(record.chair)}　　記　錄：{m(record.recorder)}",
        size=BODY_SIZE,
    )

    # ── 肆、出席人員 ──
    add_paragraph(doc, "肆、出席人員：如簽到表", size=BODY_SIZE)

    # ── 伍、會議議程 ──
    add_paragraph(doc, "伍、會議議程", size=BODY_SIZE, bold=True, space_before=Pt(4))

    # 一、主席致詞
    add_paragraph(doc, "一、主席致詞：（略）", size=BODY_SIZE)

    # 二、資源班業務報告
    add_paragraph(doc, "二、資源班業務報告", size=BODY_SIZE)
    if record.business_report:
        add_paragraph(doc, record.business_report, size=BODY_SIZE)

    section_num = 3

    # 三、前次會議決議追蹤（選用）
    if record.previous_tracking:
        add_paragraph(
            doc,
            f"{_chinese_num(section_num)}、前次會議決議追蹤",
            size=BODY_SIZE,
        )
        add_paragraph(doc, record.previous_tracking, size=BODY_SIZE)
        section_num += 1

    # 四（或三）、提案討論
    if record.proposals:
        add_paragraph(
            doc,
            f"{_chinese_num(section_num)}、提案討論",
            size=BODY_SIZE,
        )
        section_num += 1

        for i, proposal in enumerate(record.proposals, 1):
            add_paragraph(
                doc,
                f"【案由{i}】{proposal.title}",
                size=BODY_SIZE, bold=True,
            )
            add_paragraph(doc, "【說明】", size=BODY_SIZE, bold=True)
            add_paragraph(
                doc,
                m(proposal.description) if proposal.description else "（待填）",
                size=BODY_SIZE,
            )
            add_paragraph(doc, "【決議】", size=BODY_SIZE, bold=True)
            add_paragraph(doc, m(proposal.decision), size=BODY_SIZE)

    # 陸、臨時動議
    add_paragraph(
        doc,
        f"{_chinese_num(section_num)}、臨時動議",
        size=BODY_SIZE,
    )
    add_paragraph(doc, record.motions if record.motions else "無", size=BODY_SIZE)
    section_num += 1

    # 柒、散會
    add_paragraph(
        doc,
        f"{_chinese_num(section_num)}、散　　會",
        size=BODY_SIZE,
    )
    if record.time_end:
        add_paragraph(doc, record.time_end, size=BODY_SIZE)

    # ── 核章欄 ──
    add_paragraph(doc, "", size=BODY_SIZE, space_before=Pt(12))
    sign_table = doc.add_table(rows=1, cols=3)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["特教業務承辦人", "單位主管", "校長"]
    for col_idx, label in enumerate(labels):
        cell = sign_table.cell(0, col_idx)
        set_cell_text(cell, f"{label}\n\n\n\n", size=BODY_SIZE)
    add_table_borders(sign_table)

    # ════════════════════════════════════════
    # 第二頁：簽到表
    # ════════════════════════════════════════

    doc.add_page_break()

    # 簽到表標題
    add_paragraph(
        doc,
        f"南投縣{school}國民小學{year}學年度特殊教育推行委員會",
        size=TITLE_SIZE, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(2),
    )
    add_paragraph(
        doc,
        f"第 {num} 次會議記錄簽到表",
        size=TITLE_SIZE, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(6),
    )

    # 簽到表格
    sign_in_table = doc.add_table(rows=len(committee) + 1, cols=4)
    sign_in_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表頭
    headers = ["職稱", "校內職稱", "姓名", "簽到"]
    for col_idx, header in enumerate(headers):
        set_cell_text(sign_in_table.cell(0, col_idx), header, size=BODY_SIZE, bold=True)

    # 委員資料
    for row_idx, member in enumerate(committee, 1):
        title_cell = sign_in_table.cell(row_idx, 0)
        role_cell = sign_in_table.cell(row_idx, 1)
        name_cell = sign_in_table.cell(row_idx, 2)
        sign_cell = sign_in_table.cell(row_idx, 3)

        set_cell_multiline(title_cell, member.get("title", ""), size=BODY_SIZE)
        set_cell_text(role_cell, member.get("role", ""), size=BODY_SIZE)
        # 姓名中間加空格（如：林 思 遠），並套用 PII 遮蔽
        name = m(member.get("name", ""))
        spaced_name = " ".join(name) if len(name) >= 2 else name
        set_cell_text(name_cell, spaced_name, size=BODY_SIZE)
        set_cell_text(sign_cell, "", size=BODY_SIZE)

    add_table_borders(sign_in_table)

    # 儲存
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return output_path


# ── JSON mode helpers ──

def build_meeting_record(data: dict) -> MeetingRecord:
    """
    Convert a JSON input dict (from API route) into a MeetingRecord.

    Expected keys (all optional, sensible defaults applied):
      academicYear, meetingNumber, date, weekday, timeStart, timeEnd,
      location, chair, recorder, businessReport, previousTracking,
      proposals, motions, committee

    Each proposal in `proposals` should be a dict with:
      type, title, description, decision, students, refDoc
    """
    raw_proposals = data.get("proposals", [])
    proposals = []
    for p in raw_proposals:
        proposals.append(Proposal(
            type=p.get("type", "其他"),
            title=p.get("title", ""),
            description=p.get("description", ""),
            decision=p.get("decision", "（會後填入）"),
            students=p.get("students", []),
            ref_doc=p.get("refDoc", p.get("ref_doc", "")),
        ))

    committee_raw = data.get("committee", [])
    committee = committee_raw if committee_raw else DEFAULT_COMMITTEE

    return MeetingRecord(
        academic_year=data.get("academicYear", data.get("academic_year", SCHOOL_DEFAULTS["academic_year"])),
        meeting_number=data.get("meetingNumber", data.get("meeting_number", 1)),
        date=data.get("date", ""),
        weekday=data.get("weekday", ""),
        time_start=data.get("timeStart", data.get("time_start", "上午08:10")),
        time_end=data.get("timeEnd", data.get("time_end", "")),
        location=data.get("location", SCHOOL_DEFAULTS["location"]),
        chair=data.get("chair", SCHOOL_DEFAULTS["chair"]),
        recorder=data.get("recorder", SCHOOL_DEFAULTS["recorder"]),
        business_report=data.get("businessReport", data.get("business_report", "")),
        previous_tracking=data.get("previousTracking", data.get("previous_tracking", "")),
        proposals=proposals,
        motions=data.get("motions", "無"),
        committee=committee,
    )


def handle_json_mode():
    """JSON mode: read action + data from stdin, output result to stdout."""
    data = json.loads(sys.stdin.read())
    action = data.get("action")

    if action == "draft":
        # Build a Proposal from the input and call draft_proposal()
        proposal_type = data.get("proposal_type", "其他")
        students = data.get("students", [])
        ref_doc = data.get("ref_doc", "")
        title = data.get("title", proposal_type)

        proposal = Proposal(
            type=proposal_type,
            title=title,
            students=students,
            ref_doc=ref_doc,
        )

        similar = fetch_similar_meetings(proposal_type)
        description = draft_proposal(proposal, similar)

        print(json.dumps({
            "title": proposal.title,
            "description": description,
        }, ensure_ascii=False))

    elif action == "generate":
        record = build_meeting_record(data)
        names = collect_names(record)

        # Determine output path for docx
        output_dir = data.get("output_dir", os.path.expanduser("~/Downloads"))
        topics_str = "+".join(p.type for p in record.proposals) if record.proposals else "未分類"
        docx_filename = (
            f"{record.academic_year}-特推會-{record.meeting_number:02d}"
            f"-{topics_str}.docx"
        )
        docx_path = os.path.join(output_dir, docx_filename)

        docx_path = generate_docx(record, docx_path, mask_names=names)
        md_path = save_markdown(record)   # also calls update_moc() internally
        moc_path = os.path.join(OBSIDIAN_SPC_DIR, "MOC-特推會.md")
        moc_updated = os.path.exists(moc_path)

        print(json.dumps({
            "docx_path": docx_path,
            "md_path": md_path,
            "moc_updated": moc_updated,
        }, ensure_ascii=False))

    elif action == "generate-agenda":
        record = build_meeting_record(data)
        names = collect_names(record)

        # Blank out decisions for agenda version
        for p in record.proposals:
            p.decision = "（待會議決定）"

        output_dir = data.get("output_dir", os.path.expanduser("~/Downloads"))
        topics_str = "+".join(p.type for p in record.proposals) if record.proposals else "未分類"
        docx_filename = (
            f"{record.academic_year}-特推會-{record.meeting_number:02d}"
            f"-{topics_str}-議程.docx"
        )
        docx_path = os.path.join(output_dir, docx_filename)
        docx_path = generate_docx(record, docx_path, mask_names=names)

        # Save markdown with agenda status
        md_path = save_markdown(record, status="agenda-generated", mode="prep")

        result = {
            "docx_path": docx_path,
            "md_path": md_path,
            "moc_updated": True,
        }

        # Generate HTML if requested
        if data.get("pushToGitHub"):
            import sys
            sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
            from html_template import generate_agenda_html
            html_content = generate_agenda_html(record, names_to_mask=names)
            result["html_content"] = html_content

        print(json.dumps(result, ensure_ascii=False))

    else:
        print(json.dumps({"error": f"Unknown action: {action}"}), file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    if "--json" in sys.argv:
        handle_json_mode()
    else:
        print("Usage: python3 spc_meeting_core.py --json < input.json")
        print("Or use spc_meeting_cli.py for interactive mode.")
