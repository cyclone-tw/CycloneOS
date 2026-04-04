#!/usr/bin/env python3
"""
parse_docx.py — 解析 .docx 模板結構，輸出 JSON 供 LLM 判斷填入位置。
不使用佔位符，讓 LLM 自己理解模板結構。

Usage:
    python3 parse_docx.py <input.docx> [--output structure.json]
"""

import json
import sys
import os

try:
    from docx import Document
    from docx.table import Table
except ImportError:
    print("請安裝 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def parse_table(table: Table, table_index: int) -> dict:
    """解析表格結構，包含合併儲存格資訊。"""
    rows_data = []
    for ri, row in enumerate(table.rows):
        cells_data = []
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()
            # 偵測合併儲存格
            tc = cell._tc
            grid_span = tc.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
            colspan = int(grid_span.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')) if grid_span is not None else 1

            cells_data.append({
                "col": ci,
                "text": text,
                "colspan": colspan,
                "is_empty": len(text) == 0 or text in ["", " "],
            })
        rows_data.append({
            "row": ri,
            "cells": cells_data,
        })

    return {
        "type": "table",
        "index": table_index,
        "rows": len(table.rows),
        "cols": len(table.columns),
        "data": rows_data,
    }


def parse_paragraph(para, para_index: int) -> dict:
    """解析段落，包含樣式資訊。"""
    text = para.text.strip()
    style = para.style.name if para.style else "Normal"

    # 偵測勾選框
    has_checkbox = "□" in text or "☑" in text or "☐" in text

    return {
        "type": "paragraph",
        "index": para_index,
        "text": text,
        "style": style,
        "has_checkbox": has_checkbox,
        "is_empty": len(text) == 0,
        "is_heading": style.startswith("Heading") or text.startswith(("一、", "二、", "三、", "四、", "五、", "六、", "七、", "(一)", "(二)", "(三)")),
    }


def parse_docx(filepath: str) -> dict:
    """解析整份 .docx，輸出結構化 JSON。"""
    doc = Document(filepath)
    elements = []
    table_index = 0
    para_index = 0

    # python-docx 的 body 會交錯 paragraphs 和 tables
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "tbl":
            if table_index < len(doc.tables):
                elements.append(parse_table(doc.tables[table_index], table_index))
                table_index += 1

        elif tag == "p":
            if para_index < len(doc.paragraphs):
                para = doc.paragraphs[para_index]
                parsed = parse_paragraph(para, para_index)
                # 跳過空段落（減少 token 消耗）
                if not parsed["is_empty"]:
                    elements.append(parsed)
            para_index += 1

    return {
        "filename": os.path.basename(filepath),
        "total_paragraphs": len(doc.paragraphs),
        "total_tables": len(doc.tables),
        "elements": elements,
    }


def main():
    if len(sys.argv) < 2:
        print(f"Usage: {sys.argv[0]} <input.docx> [--output structure.json]", file=sys.stderr)
        sys.exit(1)

    filepath = sys.argv[1]
    output_path = None

    if "--output" in sys.argv:
        idx = sys.argv.index("--output")
        if idx + 1 < len(sys.argv):
            output_path = sys.argv[idx + 1]

    if not os.path.exists(filepath):
        print(f"File not found: {filepath}", file=sys.stderr)
        sys.exit(1)

    result = parse_docx(filepath)

    output = json.dumps(result, ensure_ascii=False, indent=2)

    if output_path:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(output)
        print(f"Structure saved to: {output_path}", file=sys.stderr)
    else:
        print(output)


if __name__ == "__main__":
    main()
