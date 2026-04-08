#!/usr/bin/env python3
"""
convert_spc_meetings.py — 批次轉換特推會會議記錄 .doc/.docx/.pdf → .md

將 110-114 學年度的特推會會議記錄轉換為結構化 Markdown，
存入 Obsidian Vault 的 02-特教業務/特推會/ 資料夾。

Usage:
    python3 convert_spc_meetings.py
    python3 convert_spc_meetings.py --dry-run   # 只列出檔案不寫入
"""

import argparse
import os
import re
import subprocess
import sys
import tempfile

from docx import Document

DRIVE_BASE = os.path.expanduser(
    "~/Library/CloudStorage/GoogleDrive-user@school.edu.tw"
    "/我的雲端硬碟/04.學校相關/01.特教業務/03.特教各項會議/01.特推會"
)

OBSIDIAN_BASE = os.path.expanduser(
    "~/Library/CloudStorage/GoogleDrive-user@gmail.com"
    "/我的雲端硬碟/Obsidian-Cyclone/02-特教業務/特推會"
)

YEAR_FOLDERS = {
    110: "110學年度特推會",
    111: "111學年度特推會",
    112: "112學年度特推會",
    113: "113學年度特推會",
    114: "114學年度特推會",
}


def find_meeting_files():
    """找出所有 110-114 學年度的會議記錄檔案（.doc/.docx/.pdf）。"""
    files = []
    for year, folder_name in YEAR_FOLDERS.items():
        folder = os.path.join(DRIVE_BASE, folder_name)
        if not os.path.isdir(folder):
            print(f"  ⚠️ 找不到資料夾: {folder_name}")
            continue
        for fname in sorted(os.listdir(folder)):
            # 只處理會議記錄，跳過簽到表、掃描檔、子資料夾等
            if not any(kw in fname for kw in ["特推會議記錄", "特推會-", "特推會("]):
                continue
            ext = os.path.splitext(fname)[1].lower()
            if ext not in (".doc", ".docx", ".pdf"):
                continue
            full = os.path.join(folder, fname)
            # 如果有 .doc 又有同名 .pdf，跳過 PDF（以 .doc 為準）
            if ext == ".pdf":
                base = os.path.splitext(fname)[0]
                if any(
                    os.path.exists(os.path.join(folder, base + e))
                    for e in [".doc", ".docx"]
                ):
                    continue
            files.append((year, fname, full))
    return files


def read_doc(path: str) -> str:
    """讀取 .doc 檔案（先用 LibreOffice 轉 .docx）。"""
    tmp = tempfile.mkdtemp(prefix="spc_")
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "docx", "--outdir", tmp, path],
        capture_output=True, timeout=60,
    )
    docx_files = [f for f in os.listdir(tmp) if f.endswith(".docx")]
    if not docx_files:
        return ""
    return read_docx(os.path.join(tmp, docx_files[0]))


def read_docx(path: str) -> str:
    """讀取 .docx 檔案，回傳純文字（保留段落和表格）。"""
    doc = Document(path)
    lines = []

    # 先收集段落
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)

    # 收集表格
    for table in doc.tables:
        lines.append("")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " / ") for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

    return "\n".join(lines)


def read_pdf(path: str) -> str:
    """讀取 PDF 檔案。"""
    import pdfplumber
    lines = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.append(text)
    return "\n".join(lines)


def parse_meeting_info(year: int, fname: str, text: str) -> dict:
    """從檔名和內容中提取會議資訊。"""
    info = {
        "academic_year": year,
        "meeting_number": 0,
        "date": "",
        "location": "",
        "chair": "",
        "recorder": "",
        "topics": "",
        "raw_text": text,
    }

    # 從檔名提取次數和主題
    m = re.search(r"第(\d+)次", fname)
    if not m:
        m = re.search(r"(\d+)第(\d+)次", fname)
        if m:
            info["meeting_number"] = int(m.group(2))
        else:
            # 110 的命名格式不同
            m = re.search(r"(\d{3})第(\d+)次", fname)
            if m:
                info["meeting_number"] = int(m.group(2))
    else:
        info["meeting_number"] = int(m.group(1))

    # 從檔名提取主題
    m = re.search(r"[（(](.+?)[）)]", fname)
    if m:
        info["topics"] = m.group(1)

    # 從內容提取日期
    m = re.search(r"(?:時間|會議時間)[：:]\s*(\d{2,3}年\d{1,2}月\d{1,2}日)", text)
    if not m:
        m = re.search(r"(\d{2,3}年\d{1,2}月\d{1,2}日)", text)
    if m:
        info["date"] = m.group(1)

    # 從內容提取地點
    m = re.search(r"(?:地點|會議地點)[：:]\s*(.+?)(?:\n|$)", text)
    if m:
        info["location"] = m.group(1).strip()

    # 從內容提取主席
    m = re.search(r"主席[：:]\s*(.+?)(?:\s{2,}|　|記錄|\n|$)", text)
    if m:
        chair = m.group(1).strip()
        # 清理「X校長YZ」→「XYZ」
        chair = re.sub(r"校長", "", chair)
        info["chair"] = chair

    # 從內容提取記錄
    m = re.search(r"記錄[者]?[：:]\s*(.+?)(?:\s{2,}|　|\n|$)", text)
    if m:
        info["recorder"] = m.group(1).strip()

    return info


def extract_sections(text: str) -> dict:
    """從全文中提取各段落結構。"""
    sections = {
        "chair_speech": "",
        "business_report": "",
        "previous_tracking": "",
        "proposals": [],
        "motions": "",
        "adjournment": "",
    }

    # 提取主席致詞
    m = re.search(r"主席致詞[：:]\s*(.+?)(?=\n[二三四五六七])", text, re.DOTALL)
    if m:
        sections["chair_speech"] = m.group(1).strip()

    # 提取業務報告
    m = re.search(
        r"(?:業務單位報告|資源班業務報告|業務報告)[：:]\s*(.+?)(?=\n[三四五六七]|提案討論)",
        text, re.DOTALL,
    )
    if m:
        sections["business_report"] = m.group(1).strip()

    # 提取前次追蹤
    m = re.search(r"前次.+?(?:追蹤|報告)[：:]\s*(.+?)(?=\n[四五六七]|提案)", text, re.DOTALL)
    if m:
        sections["previous_tracking"] = m.group(1).strip()

    # 提取提案
    proposals = []
    pattern = r"【案由[一二三四五六七八九十\d]+】\s*\n?(.+?)(?=【案由|陸、|柒、|臨時動議|散\s*會|$)"
    for pm in re.finditer(pattern, text, re.DOTALL):
        proposal_text = pm.group(0).strip()
        # 拆案由、說明、決議
        p = {"title": "", "description": "", "decision": ""}
        tm = re.search(r"【案由[^】]*】\s*\n?(.+?)(?=【說|【決|$)", proposal_text, re.DOTALL)
        if tm:
            p["title"] = tm.group(1).strip()
        dm = re.search(r"【說\s*明】\s*\n?(.+?)(?=【決|$)", proposal_text, re.DOTALL)
        if dm:
            p["description"] = dm.group(1).strip()
        rm = re.search(r"【決\s*議】\s*\n?(.+?)$", proposal_text, re.DOTALL)
        if rm:
            p["decision"] = rm.group(1).strip()
        proposals.append(p)
    sections["proposals"] = proposals

    # 臨時動議
    m = re.search(r"臨時動議[：:]\s*(.+?)(?=\n|$)", text)
    if m:
        sections["motions"] = m.group(1).strip()

    # 散會
    m = re.search(r"散\s*會[：:]\s*(.+?)(?=\n|$)", text)
    if m:
        sections["adjournment"] = m.group(1).strip()

    return sections


def build_markdown(info: dict, sections: dict) -> str:
    """組裝 Markdown 內容。"""
    # Frontmatter
    decisions = []
    for p in sections["proposals"]:
        if p["decision"]:
            # 取決議第一句作摘要
            dec = p["decision"].split("\n")[0][:80]
            decisions.append(dec)

    topics_list = [t.strip() for t in re.split(r"[、&,，]", info["topics"]) if t.strip()]

    md = "---\n"
    md += "type: 特推會會議記錄\n"
    md += f"academic_year: {info['academic_year']}\n"
    md += f"meeting_number: {info['meeting_number']}\n"
    md += f"date: \"{info['date']}\"\n"
    md += f"chair: \"{info['chair']}\"\n"
    md += f"recorder: \"{info['recorder']}\"\n"
    md += f"location: \"{info['location']}\"\n"
    md += f"topics:\n"
    for t in topics_list:
        md += f"  - {t}\n"
    md += f"decisions:\n"
    for d in decisions:
        md += f"  - \"{d}\"\n"
    md += "tags: [特推會, 會議記錄]\n"
    md += "---\n\n"

    # Title
    md += f"# {info['academic_year']}學年度 第{info['meeting_number']}次特推會會議記錄\n\n"

    # Meta
    md += f"- **日期**：{info['date']}\n"
    md += f"- **地點**：{info['location']}\n"
    md += f"- **主席**：{info['chair']}\n"
    md += f"- **記錄**：{info['recorder']}\n"
    md += f"- **主題**：{info['topics']}\n\n"
    md += "---\n\n"

    # 主席致詞
    if sections["chair_speech"] and sections["chair_speech"] != "（略）":
        md += f"## 主席致詞\n\n{sections['chair_speech']}\n\n"

    # 業務報告
    if sections["business_report"]:
        md += f"## 業務報告\n\n{sections['business_report']}\n\n"

    # 前次追蹤
    if sections["previous_tracking"]:
        md += f"## 前次會議決議追蹤\n\n{sections['previous_tracking']}\n\n"

    # 提案討論
    if sections["proposals"]:
        md += "## 提案討論\n\n"
        for i, p in enumerate(sections["proposals"], 1):
            md += f"### 案由{i}：{p['title']}\n\n"
            if p["description"]:
                md += f"**說明：**\n\n{p['description']}\n\n"
            if p["decision"]:
                md += f"**決議：**\n\n{p['decision']}\n\n"

    # 臨時動議
    if sections["motions"]:
        md += f"## 臨時動議\n\n{sections['motions']}\n\n"

    # 散會
    if sections["adjournment"]:
        md += f"## 散會\n\n{sections['adjournment']}\n\n"

    return md


def build_moc(all_records: list) -> str:
    """生成 MOC 索引。"""
    md = "---\ntype: MOC\ntags: [特推會, MOC]\n---\n\n"
    md += "# 特推會會議記錄索引\n\n"
    md += "> ○○國小特殊教育推行委員會歷次會議記錄\n\n"

    current_year = None
    for rec in sorted(all_records, key=lambda r: (r["academic_year"], r["meeting_number"])):
        if rec["academic_year"] != current_year:
            current_year = rec["academic_year"]
            md += f"\n## {current_year} 學年度\n\n"
            md += "| 次 | 日期 | 主題 | 決議摘要 |\n"
            md += "|---|---|---|---|\n"

        decisions_brief = "; ".join(rec.get("decisions_brief", []))[:60]
        link = rec["filename"].replace(".md", "")
        md += f"| [[{link}\\|第{rec['meeting_number']}次]] "
        md += f"| {rec['date']} "
        md += f"| {rec['topics']} "
        md += f"| {decisions_brief} |\n"

    md += "\n---\n\n*由 CycloneOS 自動生成*\n"
    return md


def main():
    parser = argparse.ArgumentParser(description="批次轉換特推會會議記錄")
    parser.add_argument("--dry-run", action="store_true", help="只列出檔案不寫入")
    args = parser.parse_args()

    print("🔍 掃描特推會會議記錄檔案...\n")
    files = find_meeting_files()
    print(f"找到 {len(files)} 份會議記錄\n")

    if args.dry_run:
        for year, fname, path in files:
            print(f"  {year} | {fname}")
        return

    os.makedirs(OBSIDIAN_BASE, exist_ok=True)
    all_records = []
    success = 0
    failed = 0

    for year, fname, path in files:
        ext = os.path.splitext(fname)[1].lower()
        print(f"📄 {fname}")

        try:
            if ext == ".doc":
                text = read_doc(path)
            elif ext == ".docx":
                text = read_docx(path)
            elif ext == ".pdf":
                text = read_pdf(path)
            else:
                continue

            if not text:
                print(f"  ⚠️ 無法讀取內容，跳過")
                failed += 1
                continue

            info = parse_meeting_info(year, fname, text)
            sections = extract_sections(text)

            md_content = build_markdown(info, sections)

            # 檔名格式
            num = info["meeting_number"]
            topics_short = info["topics"].replace("&", "+")[:20]
            md_filename = f"{year}-特推會-{num:02d}-{topics_short}.md"
            md_path = os.path.join(OBSIDIAN_BASE, md_filename)

            with open(md_path, "w", encoding="utf-8") as f:
                f.write(md_content)

            decisions_brief = [p["decision"].split("\n")[0][:40] for p in sections["proposals"] if p["decision"]]

            all_records.append({
                "academic_year": year,
                "meeting_number": num,
                "date": info["date"],
                "topics": info["topics"],
                "decisions_brief": decisions_brief,
                "filename": md_filename,
            })

            print(f"  ✅ → {md_filename}")
            success += 1

        except Exception as e:
            print(f"  ❌ 錯誤: {e}")
            failed += 1

    # 生成 MOC
    if all_records:
        moc_content = build_moc(all_records)
        moc_path = os.path.join(OBSIDIAN_BASE, "MOC-特推會.md")
        with open(moc_path, "w", encoding="utf-8") as f:
            f.write(moc_content)
        print(f"\n📋 MOC 索引已生成: MOC-特推會.md")

    print(f"\n{'='*50}")
    print(f"✅ 成功: {success}  ❌ 失敗: {failed}  📁 總計: {len(files)}")
    print(f"輸出位置: {OBSIDIAN_BASE}")
    print(f"{'='*50}")


if __name__ == "__main__":
    main()
